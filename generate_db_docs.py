from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_color):
    """
    fill_color: 'D9D9D9' (gray)
    """
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def generate_doc():
    doc = Document()
    
    # Styles setup for the whole document (optional but good for consistency)
    # Default font could be set here if needed
    
    # Tables data
    tables = [
        {
            "name": "Users",
            "columns": [
                ["Id", "Mã người dùng (tự tăng)", "INT", "PK"],
                ["Name", "Họ và tên người dùng", "NVARCHAR(MAX)", ""],
                ["Email", "Địa chỉ email (duy nhất)", "NVARCHAR(450)", ""],
                ["Password", "Mật khẩu đã mã hóa", "NVARCHAR(MAX)", ""],
                ["Role", "Vai trò (admin/customer)", "NVARCHAR(MAX)", ""],
                ["CreatedAt", "Thời điểm tạo tài khoản", "DATETIME", ""]
            ]
        },
        {
            "name": "Products",
            "columns": [
                ["Id", "Mã sản phẩm (tự tăng)", "INT", "PK"],
                ["Name", "Tên sản phẩm", "NVARCHAR(MAX)", ""],
                ["Description", "Mô tả chi tiết sản phẩm", "NVARCHAR(MAX)", ""],
                ["Price", "Giá bán sản phẩm", "DECIMAL(18,2)", ""],
                ["ImageUrl", "Đường dẫn hình ảnh sản phẩm", "NVARCHAR(MAX)", ""],
                ["Category", "Danh mục sản phẩm (Mice, Keyboards...)", "NVARCHAR(450)", ""],
                ["Stock", "Số lượng hàng trong kho", "INT", ""],
                ["IsPreOrder", "Trạng thái hàng đặt trước", "BIT", ""],
                ["PreOrderDate", "Ngày dự kiến có hàng đặt trước", "DATETIME", ""],
                ["IsOrderOnly", "Chỉ cho phép đặt hàng (không có sẵn)", "BIT", ""],
                ["CreatedAt", "Thời điểm thêm sản phẩm", "DATETIME", ""]
            ]
        },
        {
            "name": "Orders",
            "columns": [
                ["Id", "Mã đơn hàng (tự tăng)", "INT", "PK"],
                ["UserId", "Mã người dùng đặt hàng", "INT", "FK"],
                ["TotalPrice", "Tổng giá trị đơn hàng", "DECIMAL(18,2)", ""],
                ["Status", "Trạng thái đơn (Pending, Confirmed...)", "NVARCHAR(MAX)", ""],
                ["FullName", "Tên người nhận hàng", "NVARCHAR(MAX)", ""],
                ["PhoneNumber", "Số điện thoại nhận hàng", "NVARCHAR(MAX)", ""],
                ["ShippingAddress", "Địa chỉ giao hàng", "NVARCHAR(MAX)", ""],
                ["PaymentMethod", "Phương thức thanh toán (COD, VnPay...)", "NVARCHAR(MAX)", ""],
                ["CreatedAt", "Thời điểm tạo đơn hàng", "DATETIME", ""],
                ["ConfirmedAt", "Thời điểm xác nhận đơn", "DATETIME", ""],
                ["ShippingAt", "Thời điểm bắt đầu giao", "DATETIME", ""],
                ["DeliveredAt", "Thời điểm giao thành công", "DATETIME", ""],
                ["CancelledAt", "Thời điểm hủy đơn", "DATETIME", ""]
            ]
        },
        {
            "name": "OrderItems",
            "columns": [
                ["Id", "Mã bảng ghi (tự tăng)", "INT", "PK"],
                ["OrderId", "Mã đơn hàng cha", "INT", "FK"],
                ["ProductId", "Mã sản phẩm trong đơn", "INT", "FK"],
                ["Quantity", "Số lượng mua", "INT", ""],
                ["Price", "Giá sản phẩm tại thời điểm mua", "DECIMAL(18,2)", ""]
            ]
        },
        {
            "name": "CartItems",
            "columns": [
                ["Id", "Mã mục giỏ hàng (tự tăng)", "INT", "PK"],
                ["UserId", "Mã người dùng", "INT", "FK"],
                ["ProductId", "Mã sản phẩm", "INT", "FK"],
                ["Quantity", "Số lượng sản phẩm trong giỏ", "INT", ""]
            ]
        },
        {
            "name": "WishlistItems",
            "columns": [
                ["Id", "Mã mục yêu thích (tự tăng)", "INT", "PK"],
                ["UserId", "Mã người dùng", "INT", "FK"],
                ["ProductId", "Mã sản phẩm", "INT", "FK"],
                ["CreatedAt", "Thời điểm thêm vào yêu thích", "DATETIME", ""]
            ]
        },
        {
            "name": "Reviews",
            "columns": [
                ["Id", "Mã đánh giá (tự tăng)", "INT", "PK"],
                ["ProductId", "Mã sản phẩm được đánh giá", "INT", "FK"],
                ["UserId", "Mã người dùng đánh giá", "INT", "FK"],
                ["Rating", "Số sao (1 đến 5)", "INT", ""],
                ["Comment", "Nội dung nhận xét", "NVARCHAR(MAX)", ""],
                ["NickName", "Biệt danh người đánh giá", "NVARCHAR(MAX)", ""],
                ["CreatedAt", "Thời điểm đánh giá", "DATETIME", ""]
            ]
        },
        {
            "name": "Blogs",
            "columns": [
                ["Id", "Mã bài viết (tự tăng)", "INT", "PK"],
                ["Title", "Tiêu đề bài viết", "NVARCHAR(MAX)", ""],
                ["Content", "Nội dung chi tiết bài viết", "NVARCHAR(MAX)", ""],
                ["ImageUrl", "Hình ảnh đại diện bài viết", "NVARCHAR(MAX)", ""],
                ["Author", "Tác giả bài viết", "NVARCHAR(MAX)", ""],
                ["CreatedAt", "Thời điểm đăng tải", "DATETIME", ""]
            ]
        },
        {
            "name": "Promotions",
            "columns": [
                ["Id", "Mã khuyến mãi (tự tăng)", "INT", "PK"],
                ["Name", "Tên chương trình khuyến mãi", "NVARCHAR(200)", ""],
                ["MinOrderValue", "Giá trị đơn hàng tối thiểu", "DECIMAL(18,2)", ""],
                ["DiscountPercent", "Phần trăm giảm giá", "INT", ""],
                ["MaxDiscount", "Mức giảm tối đa", "DECIMAL(18,2)", ""],
                ["IsActive", "Trạng thái kích hoạt", "BIT", ""]
            ]
        },
        {
            "name": "Vouchers",
            "columns": [
                ["Id", "Mã Voucher (tự tăng)", "INT", "PK"],
                ["Code", "Mã code nhập vào (vd: GAME2025)", "NVARCHAR(MAX)", ""],
                ["Type", "Loại Voucher (percent/fixed)", "NVARCHAR(MAX)", ""],
                ["Value", "Giá trị giảm giá", "DECIMAL(18,2)", ""],
                ["ExpiryDate", "Ngày hết hạn sử dụng", "DATETIME", ""],
                ["MaxUsages", "Tổng số lượt dùng tối đa", "INT", ""],
                ["UsedCount", "Số lượt đã sử dụng", "INT", ""],
                ["CreatedAt", "Thời điểm tạo Voucher", "DATETIME", ""]
            ]
        },
        {
            "name": "Notifications",
            "columns": [
                ["Id", "Mã thông báo (tự tăng)", "INT", "PK"],
                ["UserId", "Mã người dùng nhận thông báo", "INT", "FK"],
                ["Message", "Nội dung thông báo", "NVARCHAR(MAX)", ""],
                ["IsRead", "Đã đọc hay chưa", "BIT", ""],
                ["CreatedAt", "Thời điểm gửi thông báo", "DATETIME", ""]
            ]
        },
        {
            "name": "ChatMessages",
            "columns": [
                ["Id", "Mã tin nhắn (tự tăng)", "INT", "PK"],
                ["SenderId", "Mã người gửi", "INT", "FK"],
                ["ReceiverId", "Mã người nhận", "INT", "FK"],
                ["Message", "Nội dung tin nhắn", "NVARCHAR(MAX)", ""],
                ["IsRead", "Đã xem hay chưa", "BIT", ""],
                ["IsAdminSender", "Xác định Admin là người gửi", "BIT", ""],
                ["CreatedAt", "Thời điểm gửi tin nhắn", "DATETIME", ""]
            ]
        },
        {
            "name": "Drivers",
            "columns": [
                ["Id", "Mã Driver (tự tăng)", "INT", "PK"],
                ["Name", "Tên phần mềm / Driver", "NVARCHAR(MAX)", ""],
                ["Description", "Mô tả chi tiết phần mềm", "NVARCHAR(MAX)", ""],
                ["DownloadUrl", "Đường dẫn tải file", "NVARCHAR(MAX)", ""],
                ["Brand", "Thương hiệu phụ kiện (vd: Logitech)", "NVARCHAR(MAX)", ""]
            ]
        },
        {
            "name": "OtpCodes",
            "columns": [
                ["Id", "Mã định danh (tự tăng)", "INT", "PK"],
                ["Email", "Email nhận mã OTP", "NVARCHAR(MAX)", ""],
                ["Code", "Mã xác thực (6 chữ số)", "NVARCHAR(6)", ""],
                ["ExpiresAt", "Thời điểm mã hết hiệu lực", "DATETIME", ""],
                ["IsUsed", "Trạng thái đã sử dụng mã", "BIT", ""],
                ["CreatedAt", "Thời điểm tạo mã", "DATETIME", ""]
            ]
        },
        {
            "name": "RefreshTokens",
            "columns": [
                ["Id", "Mã Token (tự tăng)", "INT", "PK"],
                ["Token", "Chuỗi ký tự Token", "NVARCHAR(MAX)", ""],
                ["UserId", "Mã người dùng sở hữu", "INT", "FK"],
                ["ExpiresAt", "Thời điểm hết hạn", "DATETIME", ""],
                ["IsRevoked", "Trạng thái bị thu hồi", "BIT", ""],
                ["CreatedAt", "Thời điểm tạo Token", "DATETIME", ""],
                ["RevokedAt", "Thời điểm bị thu hồi", "DATETIME", ""],
                ["ReplacedByToken", "Token mới thay thế (nếu có)", "NVARCHAR(MAX)", ""],
                ["RevokedReason", "Lý do thu hồi Token", "NVARCHAR(MAX)", ""]
            ]
        }
    ]

    for i, table_data in enumerate(tables):
        # Title above
        title_para = doc.add_paragraph()
        run = title_para.add_run(f"Bảng dữ liệu {table_data['name']}")
        run.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Create table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Header row
        hdr_cells = table.rows[0].cells
        headers = ['Tên Cột', 'Mô Tải', 'Kiểu Dữ Liệu', 'Ghi Chú']
        for idx, h in enumerate(headers):
            hdr_cells[idx].text = h
            # Bold and Center Header
            run = hdr_cells[idx].paragraphs[0].runs[0]
            run.bold = True
            hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Gray background
            set_cell_background(hdr_cells[idx], 'D9D9D9')
            
        # Add data rows
        for col_data in table_data['columns']:
            row_cells = table.add_row().cells
            for idx, val in enumerate(col_data):
                row_cells[idx].text = str(val if val is not None else "")
        
        # Caption below
        caption_para = doc.add_paragraph()
        caption_run = caption_para.add_run(f"Bảng 3.{i+1}. Bảng dữ liệu {table_data['name']}")
        caption_run.italic = True
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Spacing line between tables
        if i < len(tables) - 1:
            doc.add_paragraph()

    # Save to root of project
    output_path = r'd:\Hà Quang Huy_231A290063_WebNC\Database_Tables_Doc.docx'
    doc.save(output_path)
    print(f"File created at: {output_path}")

if __name__ == "__main__":
    generate_doc()
