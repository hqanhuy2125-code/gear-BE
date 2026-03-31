from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_color):
    """
    fill_color: 'D9D9D9' (gray)
    """
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def generate_usecase_spec_doc():
    doc = Document()
    
    # Define Use Case contents
    usecases = [
        {
            "id": "UC_01",
            "actor": "Khách vãng lai",
            "name": "Xem/Tìm kiếm sản phẩm",
            "desc": "Cho phép người dùng chưa đăng nhập có thể xem danh sách sản phẩm, chi tiết sản phẩm và tìm kiếm sản phẩm theo tên hoặc danh mục.",
            "main_flow": "1. Người dùng truy cập vào trang chủ hoặc trang sản phẩm.\n2. Hệ thống hiển thị danh sách sản phẩm.\n3. Người dùng nhập từ khóa vào ô tìm kiếm hoặc chọn danh mục.\n4. Hệ thống lọc và hiển thị các sản phẩm phù hợp.\n5. Người dùng click vào một sản phẩm để xem chi tiết.\n6. Hệ thống hiển thị thông tin chi tiết (giá, mô tả, kho...).",
            "sub_flow": "3a. Không tìm thấy sản phẩm: Hệ thống hiển thị thông báo 'Không tìm thấy sản phẩm nào'.",
            "pre_cond": "Không có.",
            "post_cond": "Người dùng nắm bắt được thông tin sản phẩm cần tìm."
        },
        {
            "id": "UC_02",
            "actor": "Khách hàng",
            "name": "Đặt hàng & Thanh toán",
            "desc": "Khách hàng thực hiện đặt mua các sản phẩm trong giỏ hàng và tiến hành thanh toán.",
            "main_flow": "1. Khách hàng truy cập vào giỏ hàng.\n2. Khách hàng kiểm tra danh sách và nhấn 'Thanh toán'.\n3. Hệ thống yêu cầu nhập thông tin giao hàng (họ tên, địa chỉ, SDT).\n4. Khách hàng chọn phương thức thanh toán (COD hoặc VNPay).\n5. Khách hàng xác nhận đặt hàng.\n6. Hệ thống tạo đơn hàng mới, cập nhật tồn kho và gửi thông báo.",
            "sub_flow": "5a. Thanh toán online thất bại: Hệ thống báo lỗi và cho phép thử lại hoặc đổi phương thức.\n6a. Hết hàng đột ngột: Hệ thống báo lỗi và dừng quy trình.",
            "pre_cond": "Khách hàng đã đăng nhập và có sản phẩm trong giỏ hàng.",
            "post_cond": "Đơn hàng mới được tạo thành công trên hệ thống."
        },
        {
            "id": "UC_03",
            "actor": "Quản trị viên",
            "name": "Quản lý sản phẩm / Blog",
            "desc": "Admin thực hiện thêm, sửa, xóa thông tin sản phẩm hoặc bài viết blog trên hệ thống.",
            "main_flow": "1. Admin truy cập trang Dashboard quản trị.\n2. Admin chọn mục 'Quản lý sản phẩm' hoặc 'Quản lý Blog'.\n3. Admin chọn một hành động (Thêm/Sửa/Xóa).\n4. Admin nhập/chỉnh sửa thông tin vào form.\n5. Admin nhấn Lưu.\n6. Hệ thống kiểm tra dữ liệu và cập nhật vào cơ sở dữ liệu.",
            "sub_flow": "4a. Dữ liệu nhập vào không hợp lệ: Hệ thống báo lỗi và yêu cầu nhập lại.",
            "pre_cond": "Tài khoản có quyền 'Quản trị viên' đã đăng nhập.",
            "post_cond": "Thông tin sản phẩm/bài viết được cập nhật thành công trên hệ thống."
        },
        {
            "id": "UC_04",
            "actor": "Shipper/Driver",
            "name": "Nhận đơn / Cập nhật giao hàng",
            "desc": "Nhân viên giao hàng tiếp nhận đơn và cập nhật trạng thái đơn hàng trong quá trình vận chuyển.",
            "main_flow": "1. Shipper đăng nhập vào hệ thống.\n2. Shipper xem danh sách các đơn hàng ở trạng thái 'Confirmed'.\n3. Shipper chọn một đơn hàng để tiếp nhận.\n4. Shipper cập nhật trạng thái sang 'Shipping' khi bắt đầu giao.\n5. Sau khi giao xong, Shipper cập nhật trạng thái đơn hàng sang 'Delivered'.\n6. Hệ thống lưu lịch sử và gửi thông báo cho khách hàng.",
            "sub_flow": "5a. Giao hàng thất bại: Shipper cập nhật lý do và chuyển trạng thái về 'Cancelled' hoặc hẹn giao lại.",
            "pre_cond": "Tài khoản có quyền 'Shipper/Driver' đã đăng nhập.",
            "post_cond": "Trạng thái đơn hàng được cập nhật chính xác trên hệ thống."
        }
    ]

    for i, uc in enumerate(usecases):
        # Header title
        title = doc.add_heading(f"Đặc tả Use Case: {uc['name']}", level=2)
        
        # Create Table
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        
        def add_row(key, val):
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = val
            # Bold for key column
            row_cells[0].paragraphs[0].runs[0].bold = True
            set_cell_background(row_cells[0], 'F2F2F2') # Light gray for first column
            
        add_row("UseCase", uc['id'])
        add_row("Tên UseCase", uc['name'])
        add_row("Actor", uc['actor'])
        add_row("Mô tả", uc['desc'])
        add_row("Dòng sự kiện chính", uc['main_flow'])
        add_row("Dòng sự kiện phụ", uc['sub_flow'])
        add_row("Điều kiện tiên quyết", uc['pre_cond'])
        add_row("Kết quả", uc['post_cond'])
        
        # Add spacing after table
        doc.add_paragraph()
        if i < len(usecases) - 1:
            doc.add_page_break()

    # Save to project root
    output_path = r'd:\Hà Quang Huy_231A290063_WebNC\BangDacTaUseCase.docx'
    doc.save(output_path)
    print(f"File created at: {output_path}")

if __name__ == "__main__":
    generate_usecase_spec_doc()
